"""
parser.py
---------
Handles raw resume ingestion:
  - Text extraction from PDF (PyMuPDF) and DOCX (python-docx)
  - Contact information extraction (name, email, phone)
  - Section segmentation (experience / education / skills / projects / certs)
  - Role, department, field-of-study, and summary extraction

This module is intentionally free of any Streamlit / UI code so it can
be unit-tested and reused independently (separation of UI and AI logic).
"""

import re
import os
from typing import Optional, List

import fitz  # PyMuPDF
import docx  # python-docx

from ai.knowledge_base import (
    SECTION_HEADERS, INDIAN_CITIES, JOB_ROLES, DEPARTMENTS, FIELDS_OF_STUDY
)

EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_REGEX = re.compile(
    r"(\+?\d{1,3}[-.\s]?)?(\(?\d{3,5}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{3,4}\b"
)
LINKEDIN_REGEX = re.compile(r"(linkedin\.com/in/[A-Za-z0-9\-_/]+)", re.IGNORECASE)
GITHUB_REGEX = re.compile(r"(github\.com/[A-Za-z0-9\-_/]+)", re.IGNORECASE)


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed."""


def extract_text(file_path: str) -> str:
    """Extract raw text from a PDF or DOCX file on disk."""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf_text(file_path)
        elif ext == ".docx":
            return _extract_docx_text(file_path)
        else:
            raise ResumeParseError(f"Unsupported file type: {ext}")
    except ResumeParseError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError(f"Failed to parse '{os.path.basename(file_path)}': {exc}") from exc


def _extract_pdf_text(file_path: str) -> str:
    text_chunks = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_chunks.append(page.get_text("text"))
    text = "\n".join(text_chunks).strip()
    if not text:
        raise ResumeParseError("No extractable text found in PDF (it may be scanned/image-based).")
    return text


def _extract_docx_text(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text out of any tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    text = "\n".join(p for p in paragraphs if p.strip())
    if not text.strip():
        raise ResumeParseError("No extractable text found in DOCX.")
    return text


def extract_email(text: str) -> Optional[str]:
    if not text:
        return None
    match = EMAIL_REGEX.search(text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    if not text:
        return None
    for line in text.split("\n"):
        match = PHONE_REGEX.search(line)
        if match:
            candidate = match.group(0).strip()
            digits = re.sub(r"\D", "", candidate)
            if 7 <= len(digits) <= 15:
                return candidate
    return None


def extract_links(text: str) -> dict:
    if not text:
        return {"linkedin": None, "github": None}
    linkedin = LINKEDIN_REGEX.search(text)
    github = GITHUB_REGEX.search(text)
    return {
        "linkedin": linkedin.group(0) if linkedin else None,
        "github": github.group(0) if github else None,
    }


def extract_name(text: str, nlp=None) -> str:
    """
    Extract candidate name. Prefers spaCy NER (PERSON entity) if a loaded
    spaCy pipeline is passed in; otherwise falls back to a heuristic that
    reads the first non-empty line that looks like a human name.
    """
    if not text:
        return "Unknown Candidate"
    
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    header_text = "\n".join(lines[:5])

    if nlp is not None:
        try:
            doc = nlp(header_text)
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    candidate = ent.text.strip()
                    if 1 <= len(candidate.split()) <= 4:
                        return candidate.title()
        except Exception:  # noqa: BLE001
            pass  # fall through to heuristic

    # Heuristic fallback: first line that is short, alphabetic-only,
    # and not an obvious header/contact line.
    skip_words = {"resume", "curriculum", "vitae", "cv", "email", "phone", "address", "linkedin", "github"}
    for line in lines[:6]:
        cleaned = line.strip(" -|:\u2022\t")
        if not cleaned or "@" in cleaned or any(ch.isdigit() for ch in cleaned):
            continue
        lower = cleaned.lower()
        if any(w in lower for w in skip_words):
            continue
        words = cleaned.split()
        if 1 <= len(words) <= 4 and all(w.replace(".", "").isalpha() for w in words):
            return cleaned.title()
    return "Unknown Candidate"


def segment_sections(text: str) -> dict:
    """
    Split resume text into labeled sections based on common header
    keywords. Returns a dict of section_name -> section_text.
    Falls back to putting everything under 'full_text' as well.
    """
    if not text:
        return {key: "" for key in SECTION_HEADERS} | {"full_text": text or ""}
    
    lines = text.split("\n")
    sections = {key: "" for key in SECTION_HEADERS}
    sections["full_text"] = text
    current = None

    header_lookup = {}
    for section, keywords in SECTION_HEADERS.items():
        for kw in keywords:
            header_lookup[kw] = section

    for line in lines:
        stripped = line.strip().lower().strip(":")
        matched_section = None
        if 2 <= len(stripped) <= 45:
            for kw, section in header_lookup.items():
                if stripped == kw or (stripped.startswith(kw) and len(stripped) < len(kw) + 6):
                    matched_section = section
                    break
        if matched_section:
            current = matched_section
            continue
        if current:
            sections[current] += line + "\n"

    return sections


def extract_location(text: str) -> str:
    """
    Extract the candidate's city/location from resume text using:
    1. "Location:" / "City:" / "Address:" prefix patterns
    2. Keyword matching against a known city names list (INDIAN_CITIES).
    Returns the city name if found, or "Not Found" otherwise.
    """
    if not text:
        return "Not Found"
    
    text_lower = text.lower()
    lines = text_lower.split("\n")

    # Priority 1: Explicit "location / city / address:" lines
    location_prefixes = [
        r"(?:location|city|current\s+city|current\s+location|address|base\s+location|place)\s*[:\-]?\s*(.+)",
        r"(?:based\s+in|located\s+in|resides?\s+in|living\s+in|settled\s+in)\s*(.+)",
    ]
    for line in lines:
        stripped = line.strip()
        for pattern in location_prefixes:
            match = re.search(pattern, stripped)
            if match:
                candidate = match.group(1).strip().strip(".,;")
                candidate_lower = candidate.lower()
                for city in INDIAN_CITIES:
                    if city in candidate_lower or candidate_lower in city:
                        return city.title()
                if candidate and len(candidate) > 1:
                    return candidate.title()

    # Priority 2: Look for city names near contact/header area (first 15 lines)
    header_lines = "\n".join(lines[:15])
    for city in INDIAN_CITIES:
        pattern = r"(?<![a-zA-Z])" + re.escape(city) + r"(?![a-zA-Z])"
        if re.search(pattern, header_lines):
            return city.title()

    # Priority 3: Search the full text for any known city name
    for city in INDIAN_CITIES:
        pattern = r"(?<![a-zA-Z])" + re.escape(city) + r"(?![a-zA-Z])"
        if re.search(pattern, text_lower):
            return city.title()

    return "Not Found"


def extract_experience_years(text: str) -> float:
    """
    Estimate total years of experience using regex patterns like
    '3 years', '2+ years of experience', '5 yrs', or a date-range
    heuristic (e.g., 2019 - 2023).
    """
    if not text or not isinstance(text, str):
        return 0.0
    
    text_lower = text.lower()

    # Range patterns like "3-5 years" or "3 to 5 years of experience"
    range_patterns = [
        r"(\d+(?:\.\d+)?)\s*(?:-|to|–)\s*(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of)?\s*experience",
        r"(\d+(?:\.\d+)?)\s*(?:-|to|–)\s*(\d+(?:\.\d+)?)\+?\s*yrs?\s*(?:of)?\s*experience",
    ]
    for pattern in range_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                return max(float(match.group(1)), float(match.group(2)))
            except ValueError:
                continue

    explicit_patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*years?\s*(?:of)?\s*experience",
        r"experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*years?",
        r"(\d+(?:\.\d+)?)\+?\s*yrs?\s*(?:of)?\s*experience",
        r"(\d+(?:\.\d+)?)\+?\s*years?\s*(?:in|at)\s+\w+",
        r"total\s+experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*years?",
        r"overall\s+experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\+?\s*years?",
    ]
    for pattern in explicit_patterns:
        match = re.search(pattern, text_lower)
        if match:
            try:
                return float(match.group(1))
            except ValueError:
                continue

    # Skip freshers
    if re.search(r"\b(fresher|entry[\s-]level|no prior experience|no work experience|recent graduate|new graduate)\b", text_lower):
        return 0.0

    # Infer from year ranges
    full_matches = re.findall(
        r"((19|20)\d{2})\s*(?:-|–|to)\s*((19|20)\d{2}|present|current)",
        text_lower,
    )
    if not full_matches:
        return 0.0
    
    import datetime
    current_year = datetime.datetime.now().year
    spans = []
    for m in full_matches:
        start_year = int(m[0])
        end_raw = m[2]
        end_year = current_year if end_raw in ("present", "current") else int(end_raw)
        if end_year >= start_year:
            spans.append((start_year, end_year))
    if spans:
        earliest = min(s[0] for s in spans)
        latest = max(s[1] for s in spans)
        return round(max(0, latest - earliest), 1)
    return 0.0


def extract_role(text: str) -> str:
    """
    Extract the candidate's current/previous job role from resume text.
    Uses keyword matching against common job roles.
    """
    if not text:
        return "Not Found"
    
    text_lower = text.lower()
    
    # Look in the summary/objective section first (first ~100 lines)
    first_lines = "\n".join(text.split("\n")[:30]).lower()
    
    # Check for explicit "Current Role:" or "Position:" patterns
    role_patterns = [
        r"(?:current\s+)?(?:role|position|title|designation)\s*[:\-]?\s*(.+)",
        r"(?:working\s+as|employed\s+as|served\s+as|worked\s+as)\s+(.+?)(?:\.|,|$|in|at)",
    ]
    for pattern in role_patterns:
        match = re.search(pattern, first_lines, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip().strip(".,;")
            for role in JOB_ROLES:
                if role in candidate.lower():
                    return role.title()
            if candidate and len(candidate) > 1:
                return candidate.title()

    # Search for role mentions in the text
    for role in JOB_ROLES:
        pattern = r"(?<![a-zA-Z])" + re.escape(role) + r"(?![a-zA-Z])"
        if re.search(pattern, text_lower):
            return role.title()

    return "Not Found"


def extract_department(text: str) -> str:
    """
    Extract the candidate's department from resume text.
    """
    if not text:
        return "Not Found"
    
    text_lower = text.lower()
    
    # Check for explicit "Department:" mentions
    dept_pattern = r"(?:department)\s*[:\-]?\s*(.+)"
    match = re.search(dept_pattern, text_lower, re.IGNORECASE)
    if match:
        candidate = match.group(1).strip().strip(".,;")
        for dept in DEPARTMENTS:
            if dept in candidate.lower():
                return dept.title()
        if candidate and len(candidate) > 1:
            return candidate.title()

    # Search for department mentions in the text
    for dept in DEPARTMENTS:
        pattern = r"(?<![a-zA-Z])" + re.escape(dept) + r"(?![a-zA-Z])"
        if re.search(pattern, text_lower):
            return dept.title()

    return "Not Found"


def extract_summary(text: str) -> str:
    """Extract the professional summary/objective from the resume."""
    if not text:
        return ""
    
    sections = segment_sections(text)
    summary_text = sections.get("summary", "").strip()
    if summary_text:
        # Clean up the summary
        lines = [l.strip() for l in summary_text.split("\n") if l.strip()]
        return " ".join(lines)[:500]
    
    # Fallback: first substantial line that isn't name/contact
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    skip_words = {"resume", "curriculum", "vitae", "cv", "email", "phone", "address", "linkedin", "github"}
    for line in lines[2:8]:  # Skip first 2 lines (usually name/contact)
        cleaned = line.strip(" -|:\u2022\t")
        if not cleaned or "@" in cleaned:
            continue
        if len(cleaned) > 30 and cleaned[0].isupper():
            return cleaned[:300]
    
    return ""


def extract_branch_specialization(text: str) -> str:
    """
    Extract the candidate's branch/specialization/field of study from education section.
    """
    if not text:
        return "Not Found"
    
    sections = segment_sections(text)
    edu_text = sections.get("education", "").lower()
    full_text = text.lower()
    
    # Search in education section first, then full text
    search_text = edu_text if edu_text else full_text
    
    for field in FIELDS_OF_STUDY:
        pattern = r"(?<![a-zA-Z])" + re.escape(field) + r"(?![a-zA-Z])"
        if re.search(pattern, search_text):
            return field.title()
    
    return "Not Found"


def extract_college_university(text: str) -> str:
    """
    Extract the candidate's college/university from education section.
    """
    if not text:
        return "Not Found"
    
    sections = segment_sections(text)
    edu_text = sections.get("education", "")
    
    if not edu_text:
        return "Not Found"
    
    # Common university/college indicators
    edu_lines = edu_text.split("\n")
    
    university_keywords = [
        r"(?:university|college|institute|institution|academy|school)\s+of\s+[\w\s]+",
        r"[\w\s]+\s+(?:university|college|institute|institution)",
        r"(?:iit|nit|iiit|vit|bits|mit|stanford|harvard|cambridge|oxford|delhi\s+university|mumbai\s+university)",
        r"(?:b\.?e?\s?college|engineering\s+college|medical\s+college)",
    ]
    
    for line in edu_lines:
        line_lower = line.lower().strip()
        for pattern in university_keywords:
            match = re.search(pattern, line_lower, re.IGNORECASE)
            if match:
                return match.group(0).strip().title()
    
    return "Not Found"


def extract_degree_name(text: str) -> str:
    """
    Extract the candidate's degree name from education section.
    """
    if not text:
        return "Not Found"
    
    from ai.knowledge_base import DEGREE_KEYWORDS, EDUCATION_RANK
    
    sections = segment_sections(text)
    edu_text = sections.get("education", "").lower()
    full_text = text.lower()
    
    search_text = edu_text if edu_text else full_text
    
    found_degrees = []
    for degree in DEGREE_KEYWORDS:
        pattern = r"(?<![a-zA-Z])" + re.escape(degree) + r"(?![a-zA-Z])"
        if re.search(pattern, search_text):
            found_degrees.append(degree)
    
    # Sort by rank descending (highest degree first)
    found_degrees = sorted(set(found_degrees), key=lambda d: -EDUCATION_RANK.get(d, 0))
    
    return ", ".join(found_degrees) if found_degrees else "Not Found"


def extract_previous_role(text: str) -> str:
    """
    Extract the candidate's previous job role (distinct from current role).
    Looks for past tense roles or roles listed under experience section.
    """
    if not text:
        return "Not Found"
    
    sections = segment_sections(text)
    exp_text = sections.get("experience", "").lower()
    
    if not exp_text:
        return "Not Found"
    
    # Look for previous role indicators
    prev_role_patterns = [
        r"(?:previous|former|past)\s+(?:role|position|job|designation)\s*[:\-]?\s*(.+)",
        r"(?:worked\s+as|was\s+(?:a|an)\s+|served\s+as)\s+(.+?)(?:\.|,|$|\n|at)",
    ]
    
    for pattern in prev_role_patterns:
        match = re.search(pattern, exp_text, re.IGNORECASE)
        if match:
            candidate = match.group(1).strip().strip(".,; ")
            for role in JOB_ROLES:
                if role in candidate.lower():
                    return role.title()
            if candidate and len(candidate) > 1:
                return candidate.title()
    
    return "Not Found"
